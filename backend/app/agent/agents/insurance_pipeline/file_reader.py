from __future__ import annotations

from pathlib import Path
import re
from typing import Dict

import fitz  # PyMuPDF
from docx import Document


class FileReader:
    def read(self, file_path: str) -> Dict[str, object]:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Không tìm thấy file: {file_path}")

        suffix = path.suffix.lower()
        if suffix == ".txt":
            raw_text = self._read_txt(path)
        elif suffix == ".docx":
            raw_text = self._read_docx(path)
        elif suffix == ".pdf":
            raw_text = self._read_pdf(path)
        else:
            raise ValueError(f"Định dạng file chưa hỗ trợ: {suffix}")

        cleaned_text = self._normalize_text(raw_text)
        return {
            "file_name": path.name,
            "file_type": suffix,
            "raw_text": raw_text,
            "cleaned_text": cleaned_text,
            "num_chars": len(cleaned_text),
        }

    def _read_txt(self, path: Path) -> str:
        return path.read_text(encoding="utf-8")

    def _read_docx(self, path: Path) -> str:
        doc = Document(str(path))
        parts = []

        # Paragraphs
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                parts.append(text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)

    def _read_pdf(self, path: Path) -> str:
        texts = []
        pdf = fitz.open(str(path))
        try:
            for page in pdf:
                text = page.get_text("text")
                if text.strip():
                    texts.append(text.strip())
        finally:
            pdf.close()
        return "\n".join(texts)

    def _normalize_text(self, text: str) -> str:
        text = text.replace("\xa0", " ")
        text = re.sub(r"\r\n|\r", "\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{2,}", "\n\n", text)
        return text.strip()


def read_text_file(file_path: str) -> Dict[str, object]:
    return FileReader().read(file_path)